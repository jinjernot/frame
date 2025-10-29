from app.routes.qs_tool.core.format.header import header
from app.routes.qs_tool.core.format.footer import footer
import json

from docx.shared import Pt
from docx.shared import RGBColor  # <<< IMPORT ADDED
from docx.enum.text import WD_BREAK

# <<< FUNCTION MODIFIED
def read_format_words_from_json(json_file):
    """
    Read bold and blue words from a JSON file.

    Parameters:
        json_file (str): The path to the JSON file.

    Returns:
        tuple: A tuple containing (list_of_bold_words, list_of_blue_words)
    """
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
        bold_words = data.get('bold_words', [])
        blue_words = data.get('blue_words', [])
        return bold_words, blue_words

def set_margins(doc):
    """
    Set document margins.

    Parameters:
        doc (docx.Document): The Word document object.
    """
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)

def set_default_font(doc):
    """
    Set default font for the document.

    Parameters:
        doc (docx.Document): The Word document object.
    """
    styles = doc.styles
    default_style = styles['Normal']
    font = default_style.font
    font.name = 'HP Forma DJR Office'
    font.size = Pt(10)

# <<< FUNCTION MODIFIED
def apply_custom_formatting(doc, bold_words, blue_words):
    """
    Apply bold font and blue color to specific words.

    Parameters:
        doc (docx.Document): The Word document object.
        bold_words (list): A list of words to be bolded.
        blue_words (list): A list of words to be colored blue.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Get the stripped text from the run
            run_text_stripped = run.text.strip()
            
            # <<< LOGIC MODIFIED (if/elif)
            # Check for bold words first
            if run_text_stripped in bold_words:
                # Add a newline before AND after the word
                run.text = "\n" + run_text_stripped + "\n"
                
                # Apply bolding after modifying the text
                run.bold = True
            
            # Check for blue words
            elif run_text_stripped in blue_words:
                # Apply blue color
                run.font.color.rgb = RGBColor(0, 0, 153)

def format_document(doc, file, imgs_path):
    """
    Apply formatting to the document.

    Parameters:
        doc (docx.Document): The Word document object.
        file (str): The path to the Word document.
        imgs_path (str): The path to the images directory.
    """
    # <<< LOGIC MODIFIED
    # Read both lists from the JSON file
    bold_words, blue_words = read_format_words_from_json('/home/garciagi/frame/app/routes/qs_tool/core/format/bold_words.json')
    #bold_words, blue_words = read_format_words_from_json('app/core/format/bold_words.json')
    
    header(doc, file)
    footer(doc, imgs_path)
    set_margins(doc)
    set_default_font(doc)
    
    # <<< LOGIC MODIFIED
    # Pass both lists to the new formatting function
    apply_custom_formatting(doc, bold_words, blue_words)

    # Apply cell spacing to all tables
    for table in doc.tables:
        table.style.paragraph_format.space_before = Pt(0)
        table.style.paragraph_format.space_after = Pt(0)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.paragraph_format.space_before = Pt(0)
                    paragraph.style.paragraph_format.space_after = Pt(0)