from docx.enum.text import WD_BREAK
import re
import pandas as pd
from docx.shared import Pt
from docx.shared import RGBColor
from app.routes.qs_tool.core.format.hr import *


def insert_paragraph(doc, df, iloc_row, iloc_column):
    """
    Insert a paragraph into both the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        iloc_row (int): The row index in the DataFrame.
        iloc_column (int): The column index in the DataFrame.
    """
    data = df.iloc[iloc_row, iloc_column]
    paragraph = doc.add_paragraph()
    paragraph.add_run(data)

def process_footnotes(doc, footnotes):
    """
    Add footnotes to the Word document with blue font color.

    Parameters:
        doc (docx.Document): The Word document object.
        footnotes (list): The list of footnotes to be added.
    """
    if not footnotes:
        return

    paragraph = doc.add_paragraph()
    for index, data in enumerate(footnotes):
        run = paragraph.add_run(data)
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnotes) - 1:
            run.add_break(WD_BREAK.LINE)

def insert_error(doc, error_message):
    """
    Insert an error message into the Word document with red font color.

    Parameters:
        doc (docx.Document): The Word document object.
        error_message (str): The error message to be added.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Error: {error_message}")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    run.add_break(WD_BREAK.LINE)

def insert_list(doc, df, start_value):
    """
    Insert a list into the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        start_value (str): The starting value for the list.
    """
    if start_value not in df.iloc[:, 1].tolist():
        insert_error(doc, f"'{start_value}' not found in DataFrame.")
        return
    
    start_index = df.index[df.iloc[:, 1] == start_value].tolist()[0]
    
    end_index = len(df)
    for i in range(start_index + 1, len(df)):
        if df.iloc[i, 1] == 'Value':
            end_index = i
            break
            
    items_df = df.iloc[start_index:end_index]
    
    non_footnotes = []
    footnotes = []

    for index, row in items_df.iterrows():
        col_a_str = str(row[0]) if not pd.isna(row[0]) else ""
        col_b_str = str(row[1]) if not pd.isna(row[1]) else ""

        if 'footnote' in col_a_str.lower():
            try:
                footnote_number = int(''.join(filter(str.isdigit, col_a_str)))
                footnotes.append(f"{footnote_number}. {col_b_str}")
            except ValueError:
                continue # 'footnote' was in text but no number
        elif col_a_str.lower().strip() == 'footnote' or col_b_str.lower().strip() == 'footnote':
            continue
        else:
            non_footnotes.append(col_b_str)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(start_value.upper()) 
    run.font.size = Pt(12)
    run.bold = True
    
    paragraph = doc.add_paragraph()

    for index, data in enumerate(non_footnotes[1:], start=1):
        pattern = re.compile(r"\[([\d,]+)\]")
        split_data = pattern.split(str(data))

        # Iterate over the split parts
        for i, text_part in enumerate(split_data):
            if i % 2 == 0:
                run = paragraph.add_run(text_part)
            else:

                sup_run = paragraph.add_run(text_part)
                sup_run.font.superscript = True
                sup_run.font.size = Pt(9) 
        
        # Add a line break if it's not the last item
        if index < len(non_footnotes) - 1:
            if 'run' in locals() and run:
                 run.add_break(WD_BREAK.LINE)
            elif 'sup_run' in locals() and sup_run:
                 sup_run.add_break(WD_BREAK.LINE)

    # Add a break after the last item
    if 'run' in locals() and run:
        run.add_break(WD_BREAK.LINE)
    elif 'sup_run' in locals() and sup_run:
        sup_run.add_break(WD_BREAK.LINE)

    process_footnotes(doc, footnotes)

    insert_horizontal_line(doc.add_paragraph(), thickness=3)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def insert_footnote(doc, df, iloc_range, iloc_column):
    """
    Insert a footnote into both the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        iloc_range (slice): The slice range for selecting footnotes.
        iloc_column (int): The column index in the DataFrame.
    """
    footnote = df.iloc[iloc_range, iloc_column].tolist()

    paragraph = doc.add_paragraph()

    for index, note in enumerate(footnote):
        run = paragraph.add_run(str(note))
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnote) - 1:
            run.add_break(WD_BREAK.LINE)