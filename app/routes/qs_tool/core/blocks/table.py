from app.routes.qs_tool.core.format.hr import *
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor, Inches, Pt
import re
from app.routes.qs_tool.core.format.table import table_column_widths

def clean_text(text):
    """Remove bullet points (• ) from text."""
    return re.sub(r"•\s*", "", text)

def process_footnotes(doc, footnotes):
    """
    Process footnotes and add them to the Word document with blue font color.
    - Replace [x] with "x." (adding a period after the number).
    - Skip unwanted values.
    """
    if not footnotes:
        return

    paragraph = doc.add_paragraph()
    pattern = re.compile(r"\[(\d+)\]")
    for index, data in enumerate(footnotes):
        data = clean_text(data)

        if "Container Name" in data or "Wireless WAN" in data:
            continue

        # Replace [x] with "x."
        formatted_text = pattern.sub(r"\1.", data)

        run = paragraph.add_run(formatted_text)
        run.font.color.rgb = RGBColor(0, 0, 153)

        if index < len(footnotes) - 1 and footnotes[index + 1].strip():
            paragraph.add_run().add_break(WD_BREAK.LINE)

def insert_table(doc, df):
    """
    Insert tables into the Word document while formatting [x] as superscripted x.
    """

    df.fillna('', inplace=True)
    df.dropna(how='all', inplace=True)

    pattern = re.compile(r"\[(\d+)\]") 

    for index, row in df.iterrows():
        if row[0] == "Table":
            page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

            table = doc.add_table(rows=1, cols=3)
            footnotes = []
            
            for i in range(index + 1, len(df)):
                
                col_a_val = str(df.iloc[i, 0]).strip().lower()
                col_b_val = str(df.iloc[i, 1]).strip().lower()
                is_footnote_marker = col_a_val in ['footnote', 'footnotes'] or col_b_val in ['footnote', 'footnotes']
                
                if df.iloc[i, 0] == "Table":
                    break
                
                elif is_footnote_marker:
                    for j in range(i + 1, len(df)):
                        next_col_a_val = str(df.iloc[j, 0]).strip().lower()
                        next_col_b_val = str(df.iloc[j, 1]).strip().lower()
                        
                        if next_col_a_val == "table":
                            break
                        if next_col_a_val in ['footnote', 'footnotes'] or next_col_b_val in ['footnote', 'footnotes']:
                            break
                        
                        col_a_data = str(df.iloc[j, 0]).strip()
                        col_b_data = str(df.iloc[j, 1]).strip()

                        if 'footnote' in col_a_data.lower() and col_b_data:
                            try:
                                footnote_number = int(''.join(filter(str.isdigit, col_a_data)))
                                formatted_text = f"[{footnote_number}] {col_b_data}"
                                footnotes.append(formatted_text)
                            except ValueError:
                                continue
                            
                        elif col_a_data.startswith('[') and col_b_data:
                             footnotes.append(f"{col_a_data} {col_b_data}")

                        elif not col_a_data and col_b_data.startswith('['):
                             footnotes.append(col_b_data)
                             
                    break 
                
                else:
                    cell_1 = table.add_row().cells[1]
                    paragraph_1 = cell_1.paragraphs[0]

                    text_data = clean_text(str(df.iloc[i, 0]))
                    split_data = pattern.split(text_data)

                    for k, text_part in enumerate(split_data):
                        run = paragraph_1.add_run(text_part)
                        if k % 2 == 0:
                            run.font.bold = True
                        else:
                            run.font.superscript = True
                            run.font.size = Pt(9)

                    cell_2 = table.rows[-1].cells[2]
                    paragraph_2 = cell_2.paragraphs[0]

                    text_data = clean_text(str(df.iloc[i, 1]))
                    split_data = pattern.split(text_data)

                    for k, text_part in enumerate(split_data):
                        run = paragraph_2.add_run(text_part)
                        if k % 2 == 1:
                            run.font.superscript = True
                            run.font.size = Pt(9)

            cell_0 = table.cell(1, 0)
            paragraph_0 = cell_0.paragraphs[0]

            text_data = clean_text(str(row[1]))
            split_data = pattern.split(text_data)

            for k, text_part in enumerate(split_data):
                run = paragraph_0.add_run(text_part)
                if k % 2 == 0:
                    run.font.bold = True
                else:
                    run.font.superscript = True
                    run.font.size = Pt(9)

            table.rows[0]._element.getparent().remove(table.rows[0]._element)

            if footnotes:
                process_footnotes(doc, footnotes)
            
            doc.add_paragraph()
            table_column_widths(table, (Inches(2), Inches(2), Inches(4)))