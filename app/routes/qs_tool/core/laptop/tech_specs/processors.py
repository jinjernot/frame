from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *

from docx.enum.text import WD_BREAK
from docx.shared import RGBColor, Pt
import pandas as pd
import re


superscript_pattern = re.compile(r"\[(\d+)\]")

def add_formatted_run(paragraph, text, bold=False, color=None, size=None):
    """
    Adds text to a paragraph, processing [n] as superscript,
    and applies optional formatting.
    """
    text = str(text)
    parts = superscript_pattern.split(text)
    
    for i, part in enumerate(parts):
        if not part: 
            continue
            
        run = paragraph.add_run(part)
        
        if i % 2 == 1:
            run.font.superscript = True
        
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size


def processors_section(doc, file):
    """Processors techspecs section"""

    try:
        # Load Excel file and check available processors sheet names.
        xls = pd.ExcelFile(file.stream, engine='openpyxl')
        sheet_name_to_use = None
        
        if "Processors" in xls.sheet_names:
            sheet_name_to_use = "Processors"
        elif "Processor" in xls.sheet_names:
            sheet_name_to_use = "Processor"
        elif "QS-Only Processors" in xls.sheet_names:
            sheet_name_to_use = "QS-Only Processors"
        else:
            raise ValueError("Sheet 'Processors' or 'Processor' or 'QS-Only Processors' not found in the Excel file.")

        # Read the sheet with header=None to preserve the original structure.
        df = pd.read_excel(file.stream, sheet_name=sheet_name_to_use, engine='openpyxl', header=None)
        df = df.fillna("")

        # Add title
        insert_title(doc, "Processors")

        # Locate the table header row that contains "Processor" and "Cores".
        header_row_idx = None
        for idx in range(len(df)):
            row_values = [str(v).strip().lower() for v in df.iloc[idx].tolist()]
            if "processor" in row_values and "cores" in row_values:
                header_row_idx = idx
                break

        if header_row_idx is None:
            raise ValueError("Could not find Processors table header row.")

        header_row = df.iloc[header_row_idx]

        # Detect optional split sub-header row (e.g., P-cores / E-cores / LP E-cores).
        sub_header_row_idx = None
        if header_row_idx + 1 < len(df):
            candidate = [str(v).strip().lower() for v in df.iloc[header_row_idx + 1].tolist()]
            if any(v in {"p-cores", "e-cores", "lp e-cores", "number of p-cores", "number of e-cores", "number of lp e-cores"} for v in candidate):
                sub_header_row_idx = header_row_idx + 1

        # Determine the start column from "Processor" and stop before owner/meta columns.
        start_col_idx = None
        header_structure = []
        meta_columns = {
            "owner",
            "owner check",
            "owner check (y/n)",
            "comments",
            "check",
            "notes",
            "remarks",
        }

        last_main_header = ""
        for col_idx in range(len(header_row)):
            main_raw = str(header_row.iloc[col_idx]).strip()
            main_lower = main_raw.lower()
            sub_raw = ""
            if sub_header_row_idx is not None:
                sub_raw = str(df.iloc[sub_header_row_idx, col_idx]).strip()
            sub_lower = sub_raw.lower()

            if start_col_idx is None and main_lower == "processor":
                start_col_idx = col_idx

            if start_col_idx is None:
                continue

            if main_lower in meta_columns or sub_lower in meta_columns:
                break

            # Fill merged/split main headers across blank main cells when sub-headers exist.
            effective_main = main_raw
            if effective_main:
                last_main_header = effective_main
            elif sub_raw and last_main_header:
                effective_main = last_main_header

            # Keep columns that carry an explicit main or sub header.
            # This keeps split columns like K/L while skipping totally empty spacer columns.
            if not main_raw and not sub_raw:
                continue

            header_structure.append({
                "col_idx": col_idx,
                "main": effective_main,
                "sub": sub_raw,
            })

        if not header_structure:
            raise ValueError("No Processors table columns found.")

        # Add optional subtitle from the row above (e.g., "Processors - AMD").
        if header_row_idx > 0:
            subtitle_row = df.iloc[header_row_idx - 1]
            subtitle_candidate = ""
            subtitle_col = start_col_idx if start_col_idx is not None else header_structure[0]["col_idx"]
            if subtitle_col < len(subtitle_row):
                subtitle_candidate = str(subtitle_row.iloc[subtitle_col]).strip()

            if subtitle_candidate and subtitle_candidate.lower() != "nan":
                subtitle_paragraph = doc.add_paragraph()
                add_formatted_run(subtitle_paragraph, subtitle_candidate, bold=True, size=Pt(12))

        # Track marker rows.
        footnotes_row_idx = None
        processor_family_row_idx = None
        for idx in range(header_row_idx + 1, len(df)):
            row_values = [str(v).strip().lower() for v in df.iloc[idx].tolist() if str(v).strip()]
            if footnotes_row_idx is None and any(v in ("footnote", "footnotes") for v in row_values):
                footnotes_row_idx = idx
            if processor_family_row_idx is None and any("processor family" in v for v in row_values):
                processor_family_row_idx = idx

        # Build data range until the first marker section.
        data_end_idx = len(df)
        if processor_family_row_idx is not None:
            data_end_idx = min(data_end_idx, processor_family_row_idx)
        if footnotes_row_idx is not None:
            data_end_idx = min(data_end_idx, footnotes_row_idx)

        selected_cols = [h["col_idx"] for h in header_structure]
        data_start_idx = header_row_idx + (2 if sub_header_row_idx is not None else 1)
        data_df = df.iloc[data_start_idx:data_end_idx, selected_cols].copy()

        # Create processors table with one or two header rows.
        has_subheaders = any(str(h["sub"]).strip() for h in header_structure)
        table = doc.add_table(rows=2 if has_subheaders else 1, cols=len(header_structure))
        table.autofit = True

        if has_subheaders:
            main_cells = table.rows[0].cells
            sub_cells = table.rows[1].cells

            for i, header_info in enumerate(header_structure):
                main_cells[i].text = ""
                sub_cells[i].text = ""
                add_formatted_run(main_cells[i].paragraphs[0], header_info["main"], bold=True)
                if header_info["sub"]:
                    add_formatted_run(sub_cells[i].paragraphs[0], header_info["sub"], bold=True)

            # Horizontal merge for grouped main headers (e.g., Max Turbo Frequency over J/K/L).
            group_start = 0
            while group_start < len(header_structure):
                group_main = header_structure[group_start]["main"]
                group_end = group_start
                while group_end + 1 < len(header_structure) and header_structure[group_end + 1]["main"] == group_main:
                    group_end += 1

                if group_end > group_start:
                    merged = main_cells[group_start].merge(main_cells[group_end])
                    merged.text = ""
                    add_formatted_run(merged.paragraphs[0], group_main, bold=True)

                group_start = group_end + 1

            # Vertical merge for non-split columns.
            for i, header_info in enumerate(header_structure):
                if not str(header_info["sub"]).strip():
                    sub_cells[i].merge(main_cells[i])
        else:
            header_cells = table.rows[0].cells
            for i, header_info in enumerate(header_structure):
                header_cells[i].text = ""
                add_formatted_run(header_cells[i].paragraphs[0], header_info["main"], bold=True)

        for _, data_row in data_df.iterrows():
            row_values = [str(v).strip() for v in data_row.tolist()]

            # Skip empty rows in the body.
            if not any(row_values):
                continue

            # Defensive guard if markers are embedded in selected columns.
            row_values_lower = [v.lower() for v in row_values]
            if any("footnote" in v for v in row_values_lower):
                break
            if any("processor family" in v for v in row_values_lower):
                break

            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_values):
                row_cells[i].text = ""
                add_formatted_run(row_cells[i].paragraphs[0], cell_text)

        # Render "Processor Family" list after the table.
        if processor_family_row_idx is not None:
            subtitle_paragraph = doc.add_paragraph()
            add_formatted_run(subtitle_paragraph, "Processor Family", bold=True, size=Pt(12))

            processor_family_values = []
            seen_processor_family_values = set()
            values_end_idx = footnotes_row_idx if footnotes_row_idx is not None else len(df)

            for idx in range(processor_family_row_idx + 1, values_end_idx):
                row_values = [str(v).strip() for v in df.iloc[idx].tolist() if str(v).strip() and str(v).strip().lower() != "nan"]
                if not row_values:
                    continue

                value = row_values[0]
                if value.lower() in ("footnote", "footnotes"):
                    continue
                if value not in seen_processor_family_values:
                    seen_processor_family_values.add(value)
                    processor_family_values.append(value)

            if processor_family_values:
                values_paragraph = doc.add_paragraph()
                for idx, value in enumerate(processor_family_values):
                    add_formatted_run(values_paragraph, value)
                    if idx < len(processor_family_values) - 1:
                        values_paragraph.add_run().add_break(WD_BREAK.LINE)

        # Render footnotes after Processor Family.
        if footnotes_row_idx is not None:
            doc.add_paragraph()
            footnotes_data = df.iloc[footnotes_row_idx + 1:].dropna(how='all')

            formatted_footnotes = []
            seen_footnotes = set()
            footnote_color = RGBColor(0, 0, 153)

            for _, row in footnotes_data.iterrows():
                if len(row) < 2:
                    continue

                col_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                col_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else ""

                if not col_a and not col_b:
                    continue
                if col_a.lower() in ("footnote", "footnotes"):
                    continue

                if "footnote" in col_a.lower() and col_b:
                    footnote_number_str = ''.join(filter(str.isdigit, col_a))
                    if footnote_number_str:
                        formatted_text = f"{int(footnote_number_str)}. {col_b}"
                        if formatted_text not in seen_footnotes:
                            seen_footnotes.add(formatted_text)
                            formatted_footnotes.append(formatted_text)

            if formatted_footnotes:
                footnote_paragraph = doc.add_paragraph()
                for index, footnote_text in enumerate(formatted_footnotes):
                    add_formatted_run(footnote_paragraph, footnote_text, color=footnote_color)
                    if index < len(formatted_footnotes) - 1:
                        footnote_paragraph.add_run().add_break(WD_BREAK.LINE)


        # Insert Horizontal Line
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        # Insert Page Break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        error_msg = f"An error occurred: {e}"
        print(error_msg)  # Log error to console

        return str(e)