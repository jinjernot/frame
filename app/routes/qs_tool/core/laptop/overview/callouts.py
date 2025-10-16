from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from app.routes.qs_tool.core.format.table import table_column_widths
from docx.shared import Pt, Inches
import pandas as pd
import requests
from io import BytesIO
from app.routes.qs_tool.core.format.hr import insert_horizontal_line
import os

def download_image(url):
    """Download image from URL and return the image data."""
    # Check if URL is a non-empty string
    if not isinstance(url, str) or not url.strip():
        print("Invalid URL provided for image download.")
        return None
        
    try:
        # Make the request to download the image
        response = requests.get(url, timeout=10)
        # Check if the request was successful
        if response.status_code == 200:
            # Return the image content in a BytesIO object
            return BytesIO(response.content)
        else:
            # Print an error if the image could not be downloaded
            print(f"Failed to download image from {url}. Status code: {response.status_code}")
            return None
    except requests.exceptions.RequestException as e:
        # Print an error if the request fails
        print(f"Error downloading image from {url}: {e}")
        return None

def get_temp_filename(counter, suffix=".png"):
    """Generate a fixed temporary file name with a three-digit counter."""
    return f"image{counter:03d}{suffix}"


def callout_section(doc, file, prod_name, df):
    """Add Callout Section"""
    print("inside callout section...")
    # Add the product name
    prodname_paragraph = doc.add_paragraph()
    run = prodname_paragraph.add_run(prod_name)
    run.font.size = Pt(14)
    run.bold = True
    prodname_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    print("Added product name...")
    
    # Read the data from the 'Callouts' sheet in the Excel file
    #df = pd.read_excel(file, sheet_name='Callouts')
    print("Read Callouts sheet...")
    df = pd.read_excel(file.stream, sheet_name='Callouts', engine='openpyxl')

    # Set the target directory for saving images
    target_directory = '/home/garciagi/qs'
    #target_directory = '.'

    # Get image URLs from the DataFrame
    img_url1 = df.iloc[4, 0]
    img_url2 = df.iloc[11, 0]

    # Initialize image counter
    img_counter = 1

    # Download images
    img_data1 = download_image(img_url1)
    img_data2 = download_image(img_url2)

    # --- Image 1 Handling ---
    if img_data1:
        img_filename1 = get_temp_filename(img_counter)
        img_filepath1 = os.path.join(target_directory, img_filename1)
        
        try:
            # Save image to the specified directory
            with open(img_filepath1, "wb") as img_file1:
                img_file1.write(img_data1.getvalue())
            print(f"Saved image 1 to {img_filepath1}")

            # Add image to the document
            paragraph_with_image = doc.add_paragraph()
            run = paragraph_with_image.add_run()
            run.add_picture(img_filepath1, width=Inches(5))
            paragraph_with_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error saving or adding image 1: {e}")
    else:
        print("Skipping image 1 because download failed or URL was invalid.")
    
    img_counter += 1

    # Add Front subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Front")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Table "Front"
    data_range = df.iloc[4:10, 1:5].dropna(how='all')
    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                cell.text = str(int(value)) if isinstance(value, (int, float)) else str(value)
    
    # Insert HR
    insert_horizontal_line(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # --- Image 2 Handling ---
    if img_data2:
        img_filename2 = get_temp_filename(img_counter)
        img_filepath2 = os.path.join(target_directory, img_filename2)
        
        try:
            # Save image to the specified directory
            with open(img_filepath2, "wb") as img_file2:
                img_file2.write(img_data2.getvalue())
            print(f"Saved image 2 to {img_filepath2}")

            # Add image to the document
            paragraph_with_image = doc.add_paragraph()
            run = paragraph_with_image.add_run()
            run.add_picture(img_filepath2, width=Inches(5))
            paragraph_with_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error saving or adding image 2: {e}")
    else:
        print("Skipping image 2 because download failed or URL was invalid.")

    # Add Sides subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Sides")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("Added sides subtitle...")
    
    # Add table 'Sides'
    data_range = df.iloc[11:23, 1:5].dropna(how='all')
    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                cell.text = str(int(value)) if isinstance(value, (int, float)) else str(value)

    # Insert HR
    insert_horizontal_line(doc.add_paragraph(), thickness=3)

    doc.add_page_break()
    section = doc.sections[-1]
    section.start_type = WD_SECTION.CONTINUOUS

