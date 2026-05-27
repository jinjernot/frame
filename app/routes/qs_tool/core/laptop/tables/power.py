from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.blocks.table import *
from app.routes.qs_tool.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def unbold_power_labels(doc, table_count_before):
    """In power tables: bold only section header rows (label with no value).
    Any label row that has a value in col 2 should not be bold."""
    for table in doc.tables[table_count_before:]:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 3:
                continue
            label_text = cells[1].text.strip()
            value_text = cells[2].text.strip()
            # Data label row: has both a label and a value → remove bold from label
            if label_text and value_text:
                for paragraph in cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = False


def power_section(doc, file):
    """Power QS Only Section"""

    try:
        # Load xlsx
<<<<<<< HEAD
        df = pd.read_excel(file.stream, sheet_name='QS-Only Power', engine='openpyxl')
=======
        df = pd.read_excel(file.stream, sheet_name='QS-Only Power', engine='openpyxl', header=None)
>>>>>>> d4eaf05 (Saving local changes before syncing with remote)

        # Add title: Power
        insert_title(doc, "POWER")
        
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Power supply availability may vary by country.")
        run = paragraph.add_run("Battery is internal and replaceable by customer. Serviceable by warranty.  ")
        run.font.color.rgb = RGBColor(0, 0, 153) 
        paragraph.add_run().add_break(WD_BREAK.LINE)

<<<<<<< HEAD
=======
        # Add subtitle from the marker row pattern:
        # "Container Name | Value" followed by "Power | <subtitle>" (e.g. Battery).
        subtitle_added = False
        for i in range(1, len(df)):
            prev_col0 = str(df.iloc[i - 1, 0]).strip().lower() if not pd.isna(df.iloc[i - 1, 0]) else ""
            prev_col1 = str(df.iloc[i - 1, 1]).strip().lower() if not pd.isna(df.iloc[i - 1, 1]) else ""
            curr_col0 = str(df.iloc[i, 0]).strip().lower() if not pd.isna(df.iloc[i, 0]) else ""
            curr_col1 = str(df.iloc[i, 1]).strip() if not pd.isna(df.iloc[i, 1]) else ""

            if (
                prev_col0 == "container name"
                and prev_col1 == "value"
                and curr_col0 == "power"
                and curr_col1
            ):
                insert_subtitle(doc, df, i, 1)
                subtitle_added = True
                break

        # Fallback if the expected marker format differs but a Power row exists.
        if not subtitle_added:
            for i in range(len(df)):
                curr_col0 = str(df.iloc[i, 0]).strip().lower() if not pd.isna(df.iloc[i, 0]) else ""
                curr_col1 = str(df.iloc[i, 1]).strip() if not pd.isna(df.iloc[i, 1]) else ""
                if curr_col0 == "power" and curr_col1 and curr_col1.lower() != "value":
                    insert_subtitle(doc, df, i, 1)
                    break

>>>>>>> d4eaf05 (Saving local changes before syncing with remote)
        # Add table
        table_count_before = len(doc.tables)
        insert_table(doc, df)
        unbold_power_labels(doc, table_count_before)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)