from app.routes.qs_tool.core.format.table import table_column_widths
from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *
from docx.shared import Inches, Pt, RGBColor

from docx.enum.text import WD_BREAK
import pandas as pd

def options_section(doc, file):
    """Options QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Options', engine='openpyxl', header=None)

        # Add title: Options
        insert_title(doc, "OPTIONS")
        
        start_col_idx = 0
        end_col_idx = 2
        header_row_idx = 2  # Row 3 in Excel (Container Name, Description, Part Number)
        start_row_idx = 3   # Start data from row 4
        end_row_idx = 299

        # Get the header row first
        header_data = df.iloc[header_row_idx, start_col_idx:end_col_idx+1].tolist()
        
        # Get the data rows
        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=1, cols=num_cols)
        table.autofit = False

        # Add header row first
        header_cells = table.rows[0].cells
        for col_idx in range(len(header_data)):
            header_value = header_data[col_idx]
            if not pd.isna(header_value) and str(header_value).strip():
                run = header_cells[col_idx].paragraphs[0].add_run(str(header_value))
                run.font.bold = True

        for row_idx in range(num_rows):
            row = data_range.iloc[row_idx]
            if row.isna().all(): 
                break
            row_cells = table.add_row().cells

            for col_idx in range(num_cols):
                value = row.iloc[col_idx]
                
                if not pd.isna(value):
                    run = row_cells[col_idx].paragraphs[0].add_run(str(value))
                    
                    if col_idx == 0:  # Bold first column (section names and item categories)
                        run.font.bold = True
                            
        table_column_widths(table, (Inches(2), Inches(4), Inches(2)))

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)

