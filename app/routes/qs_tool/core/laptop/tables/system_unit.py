from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK
import pandas as pd
import re

def system_unit_section(doc, file):
    """System Unit table"""
    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only System Unit', engine='openpyxl')

        # Add title: SYSTEM UNIT
        insert_title(doc, "SYSTEM UNIT")

        start_col_idx = 0
        end_col_idx = 1
        start_row_idx = 4
        end_row_idx = 42  # Default end row if no marker is found

        # --- UPDATED LOGIC TO FIND END OF TABLE ---
        # Search for 'Footnotes' (plural) or 'Footnote' (singular) in the first column
        col_a_search = df.iloc[:, 0].astype(str).str.strip()
        footnotes_plural_idx = col_a_search[col_a_search == 'Footnotes'].index.tolist()
        footnotes_singular_idx = col_a_search[col_a_search == 'Footnote'].index.tolist()

        all_footnote_marker_indices = sorted(footnotes_plural_idx + footnotes_singular_idx)

        first_footnote_marker_index = None
        if all_footnote_marker_indices:
            # Get the first marker found
            first_footnote_marker_index = all_footnote_marker_indices[0]
            # Stop the table data one row *before* the marker
            end_row_idx = first_footnote_marker_index - 1
        # --- END UPDATED LOGIC ---

        # Extract only the necessary rows for the table
        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Define column widths
        column_widths = (Inches(3), Inches(5))
        table_column_widths(table, column_widths)

        # This pattern handles the superscript [1] in the main table
        pattern = re.compile(r"\[(\d+)\]")  # Match [x] where x is a number

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                cell = table.cell(row_idx, col_idx)
                
                if not pd.isna(value):
                    data = str(value)
                    split_data = pattern.split(data)  # Splitting text while keeping numbers
                    paragraph = cell.paragraphs[0]  # Get the first paragraph in the cell

                    for i, text_part in enumerate(split_data):
                        run = paragraph.add_run(text_part)
                        if i % 2 == 1:  # If it's a matched number
                            run.font.superscript = True  # Apply superscript
                            run.font.size = Pt(9)  # Adjust font size

        # Bold the first column
        for row in table.rows:
            if row.cells[0].paragraphs and row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].font.bold = True

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                                
        doc.add_paragraph()

        # --- UPDATED FOOTNOTE PROCESSING LOGIC ---
        # Process footnotes if a marker was found
        if first_footnote_marker_index is not None:
            # Start processing from the row *after* the marker
            footnotes_data = df.iloc[first_footnote_marker_index + 1:].dropna(how='all')
            
            formatted_footnotes = []

            for _, row in footnotes_data.iterrows():
                # Ensure row has at least 2 columns and data is not NaN
                if len(row) < 2 or pd.isna(row.iloc[0]) or pd.isna(row.iloc[1]):
                    continue
                    
                col_a = str(row.iloc[0]).lower().strip()
                col_b = str(row.iloc[1]).strip()

                # Skip known irrelevant rows
                if "container name" in col_a or "wireless wan" in col_a:
                    continue

                # Skip header row if it got included (e.g., 'Footnotes' or 'Footnote')
                if col_a == 'footnotes' or col_a == 'footnote':
                    continue
                
                # NEW LOGIC: Checks for "footnote1", "footnote2", etc.
                if 'footnote' in col_a:
                    try:
                        # Extract number (e.g., from "footnote1")
                        footnote_number = int(''.join(filter(str.isdigit, col_a)))
                        # Format as "1. Text from Column B"
                        formatted_text = f"{footnote_number}. {col_b}"
                        formatted_footnotes.append(formatted_text)
                    except ValueError:
                        # 'footnote' was in the text but no number was found
                        continue
                
                # Fallback for old format: [1] in Col A, Text in Col B
                elif pattern.match(col_a):
                    try:
                        footnote_number = pattern.match(col_a).group(1)
                        formatted_text = f"{footnote_number}. {col_b}"
                        formatted_footnotes.append(formatted_text)
                    except (ValueError, AttributeError):
                        continue

            # Add the processed footnotes to the document
            if formatted_footnotes:
                paragraph = doc.add_paragraph()
                for index, data in enumerate(formatted_footnotes):
                    run = paragraph.add_run(data)
                    run.font.color.rgb = RGBColor(0, 0, 153)  # Set font color to blue
                    
                    # Add a line break if it's not the last item
                    if index < len(formatted_footnotes) - 1:
                        run.add_break(WD_BREAK.LINE)
        # --- END FOOTNOTE PROCESSING ---

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)

def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width