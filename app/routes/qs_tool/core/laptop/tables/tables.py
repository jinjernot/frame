from app.routes.qs_tool.core.laptop.tables.system_unit import system_unit_section
from app.routes.qs_tool.core.laptop.tables.displays import displays_section
from app.routes.qs_tool.core.laptop.tables.audio import audio_section
from app.routes.qs_tool.core.laptop.tables.fingerprint import fingerprint_section
from app.routes.qs_tool.core.laptop.tables.storage import storage_section
from app.routes.qs_tool.core.laptop.tables.network import network_section
from app.routes.qs_tool.core.laptop.tables.power import power_section
from app.routes.qs_tool.core.laptop.tables.options import options_section
from app.routes.qs_tool.core.laptop.tables.change_log import change_log_section
from docx.shared import RGBColor


def _is_document_issue(error_message):
    if not isinstance(error_message, str):
        return False

    message = error_message.lower()
    document_issue_markers = [
        "worksheet named",
        "no sheet named",
        "sheet",
        "sheets",
    ]
    return any(marker in message for marker in document_issue_markers)


def _render_section_error(doc, section_name, error_message):
    if not isinstance(error_message, str) or not error_message.strip():
        return

    if not _is_document_issue(error_message):
        return

    para = doc.add_paragraph()
    run = para.add_run(f"An error occurred in {section_name}: {error_message}")
    run.font.color.rgb = RGBColor(255, 0, 0)


def table_section(doc, file):
    """Table Secion"""

    # Table sections
    sections = [
        ("System Unit", system_unit_section),
        ("QS-Only Displays", displays_section),
        ("QS-Only Storage", storage_section),
        ("QS-Only Network", network_section),
        ("QS-Only Power", power_section),
        ("QS-Only Audio", audio_section),
        ("QS-Only Fingerprint", fingerprint_section),
        ("QS-Only Options", options_section),
        ("Changelog", change_log_section),
    ]

    for section_name, section_func in sections:
        section_error = section_func(doc, file)
        _render_section_error(doc, section_name, section_error)

